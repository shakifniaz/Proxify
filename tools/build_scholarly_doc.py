from pathlib import Path
import math
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\shaki\Documents\Code\WD\proxify")
OUT = ROOT / "output" / "documents"
ASSETS = ROOT / "storage" / "logs" / "scholarly_doc_assets"
OUT.mkdir(parents=True, exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "Scholarly_Project_Documentation_With_Figure_Descriptions.docx"

DARK = "10211B"
INK = "1E2924"
GREEN = "09B884"
GREEN_DARK = "04795A"
MINT = "8BED9A"
MINT_LIGHT = "BDF8C8"
PAPER = "F7FBF8"
BORDER = "D9E8DF"
MUTED = "5B6B7F"
FONT = "Aptos"
MONO = "Consolas"


def clean(text):
    if text is None:
        return ""
    for a, b in {
        "\u2014": "-",
        "\u2013": "-",
        "\u2212": "-",
        "\u2010": "-",
        "\u2011": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
        "\u2192": "->",
        "\u00a0": " ",
    }.items():
        text = text.replace(a, b)
    return text


def set_font(run, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths=None, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    if widths:
        set_widths(table, widths)
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            border(cell, GREEN_DARK if header and r == 0 else BORDER)
            margins(cell)
            if header and r == 0:
                shade(cell, GREEN_DARK)
            elif r % 2 == 0:
                shade(cell, "FBFEFC")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_font(run)
                    run.font.size = Pt(9.2)
                    if header and r == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        run.font.color.rgb = RGBColor.from_string(INK)


def p(doc, text="", bold=False, color=None, size=None, align=None, after=6, before=0):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.12
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(clean(text))
        set_font(run)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if size:
            run.font.size = Pt(size)
    return paragraph


def h(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(15 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(clean(text))
    set_font(run)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN_DARK if level <= 2 else DARK)
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11.5)
    return paragraph


def caption(doc, text):
    paragraph = p(doc, text, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    for run in paragraph.runs:
        run.italic = True


def figure_description(doc, text):
    return p(doc, text, size=9.2, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10)


def bullets(doc, items, numbered=False):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(clean(item))
        set_font(run)
        run.font.size = Pt(10.2)
        run.font.color.rgb = RGBColor.from_string(INK)


def code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.15)
    run = paragraph.add_run(clean(text))
    set_font(run, MONO)
    run.font.size = Pt(8.3)
    run.font.color.rgb = RGBColor.from_string(DARK)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF7F1")
    paragraph._p.get_or_add_pPr().append(shd)


def kv(doc, rows, widths=(1.8, 4.7)):
    table = doc.add_table(rows=len(rows), cols=2)
    for i, (key, value) in enumerate(rows):
        table.cell(i, 0).text = clean(key)
        table.cell(i, 1).text = clean(value)
    style_table(table, widths=widths, header=False)
    for row in table.rows:
        shade(row.cells[0], "EAF8EF")
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(GREEN_DARK)
    return table


def table(doc, rows, widths):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            t.cell(i, j).text = clean(value)
    style_table(t, widths=widths)
    return t


def pil_font(size, bold=False):
    choices = [
        r"C:\Windows\Fonts\aptos-bold.ttf" if bold else r"C:\Windows\Fonts\aptos.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for choice in choices:
        if Path(choice).exists():
            return ImageFont.truetype(choice, size)
    return ImageFont.load_default()


def center(draw, box, text, font_obj, fill="#10211B"):
    lines = text.split("\n")
    dims = [draw.textbbox((0, 0), line, font=font_obj) for line in lines]
    widths = [d[2] - d[0] for d in dims]
    heights = [d[3] - d[1] for d in dims]
    y = box[1] + ((box[3] - box[1]) - (sum(heights) + (len(lines) - 1) * 6)) / 2
    for line, width, height in zip(lines, widths, heights):
        x = box[0] + ((box[2] - box[0]) - width) / 2
        draw.text((x, y), line, font=font_obj, fill=fill)
        y += height + 6


def rounded(draw, box, radius, fill, outline, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_logo(path):
    size = 700
    img = Image.new("RGBA", (size, size), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img, "RGBA")

    def s(points):
        return [(int(x * size / 500), int(y * size / 500)) for x, y in points]

    draw.polygon(s([(250, 467.52), (8.34, 391.87), (8.34, 103.56), (250, 179.21)]), fill=(139, 237, 154, 175))
    draw.polygon(s([(250, 467.52), (66.43, 320.79), (66.43, 32.48), (250, 179.21)]), fill=(4, 121, 90, 255))
    draw.polygon(s([(250, 467.52), (491.66, 391.87), (491.66, 103.56), (250, 179.21)]), fill=(189, 248, 200, 185))
    draw.polygon(s([(250, 467.52), (433.57, 320.79), (433.57, 32.48), (250, 179.21)]), fill=(9, 184, 132, 255))
    img.save(path)


def draw_architecture(path):
    img = Image.new("RGB", (1400, 780), "#F7FBF8")
    draw = ImageDraw.Draw(img)
    draw.text((45, 35), "Scholarly System Architecture", font=pil_font(34, True), fill="#10211B")
    boxes = [
        ((70, 150, 360, 300), "User Interface\nVue 3 + Inertia"),
        ((555, 150, 845, 300), "Laravel 12\nControllers + Services"),
        ((1040, 150, 1330, 300), "Data Layer\nMySQL + Firebase"),
        ((130, 450, 420, 600), "Auth\nLaravel + Firebase"),
        ((555, 450, 845, 600), "Domain Modules\nRoutine, Proxy, Exams"),
        ((980, 450, 1270, 600), "Integrations\nWhatsApp Cloud API"),
    ]
    for box, label in boxes:
        rounded(draw, box, 24, "#FFFFFF", "#09B884", 4)
        center(draw, box, label, pil_font(22, True))
    arrows = [((360, 225), (555, 225)), ((845, 225), (1040, 225)), ((275, 300), (275, 450)), ((700, 300), (700, 450)), ((845, 525), (980, 525)), ((420, 525), (555, 525))]
    for start, end in arrows:
        draw.line([start, end], fill="#04795A", width=5)
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        p1 = (end[0] - 18 * math.cos(angle) - 10 * math.sin(angle), end[1] - 18 * math.sin(angle) + 10 * math.cos(angle))
        p2 = (end[0] - 18 * math.cos(angle) + 10 * math.sin(angle), end[1] - 18 * math.sin(angle) - 10 * math.cos(angle))
        draw.polygon([end, p1, p2], fill="#04795A")
    draw.text((70, 675), "Laravel serves the Vue/Inertia frontend. MySQL stores administrative records. Firebase stores live classroom streams.", font=pil_font(18), fill="#5B6B7F")
    img.save(path)


def draw_usecase(path):
    img = Image.new("RGB", (1400, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((45, 35), "Scholarly Use Case Overview", font=pil_font(34, True), fill="#10211B")
    rounded(draw, (320, 110, 1100, 820), 30, "#F7FBF8", "#BDF8C8", 4)
    actors = [("Admin", (100, 185)), ("Teacher", (100, 420)), ("Student", (100, 650))]
    for name, (x, y) in actors:
        draw.ellipse((x + 35, y, x + 85, y + 50), fill="#09B884")
        draw.line((x + 60, y + 50, x + 60, y + 125), fill="#10211B", width=4)
        draw.line((x + 20, y + 80, x + 100, y + 80), fill="#10211B", width=4)
        draw.line((x + 60, y + 125, x + 25, y + 175), fill="#10211B", width=4)
        draw.line((x + 60, y + 125, x + 95, y + 175), fill="#10211B", width=4)
        draw.text((x + 15, y + 185), name, font=pil_font(21, True), fill="#10211B")
    cases = [
        ("Manage institution", 520, 150), ("Build routines", 780, 150),
        ("Approve proxy plans", 520, 275), ("Publish notices", 780, 275),
        ("Post classroom work", 520, 430), ("Request leave", 780, 430),
        ("View schedules", 520, 585), ("Submit assignments", 780, 585),
        ("Receive updates", 650, 720),
    ]
    for label, x, y in cases:
        draw.ellipse((x - 140, y - 38, x + 140, y + 38), fill="#FFFFFF", outline="#09B884", width=3)
        center(draw, (x - 140, y - 38, x + 140, y + 38), label, pil_font(16))
    for start, end in [((200, 280), (520, 150)), ((200, 280), (780, 150)), ((200, 280), (520, 275)), ((200, 280), (780, 275)), ((200, 515), (520, 430)), ((200, 515), (780, 430)), ((200, 515), (650, 720)), ((200, 745), (520, 585)), ((200, 745), (780, 585)), ((200, 745), (650, 720))]:
        draw.line([start, end], fill="#9BBBAA", width=2)
    img.save(path)


def draw_erd(path):
    img = Image.new("RGB", (1500, 980), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((45, 35), "Core Entity Relationship Diagram", font=pil_font(34, True), fill="#10211B")
    entities = {
        "users": ["id PK", "institution_id FK", "teacher_profile_id FK", "class_section_id FK", "role", "settings"],
        "institutions": ["id PK", "owner_user_id FK", "name", "academic_year", "settings"],
        "teacher_profiles": ["id PK", "institution_id FK", "user_id FK", "name", "whatsapp_number"],
        "class_sections": ["id PK", "institution_id FK", "class_teacher_profile_id FK", "class_name", "subjects JSON"],
        "routines": ["id PK", "institution_id FK", "classes JSON", "generated_grid JSON", "teacher_schedule JSON"],
        "proxy_runs": ["id PK", "routine_id FK", "absent_teachers JSON", "assignments JSON", "status"],
        "notices": ["id PK", "institution_id FK", "user_id FK", "board", "urgency"],
        "leave_requests": ["id PK", "routine_id FK", "teacher_id", "status", "proxy_relevant"],
        "exam_schedules": ["id PK", "routine_id FK", "exam_grid JSON", "status"],
        "proxy_message_logs": ["id PK", "proxy_run_id FK", "teacher_profile_id FK", "status", "provider_message_id"],
    }
    positions = {
        "institutions": (60, 130), "users": (430, 130), "teacher_profiles": (800, 130), "class_sections": (1160, 130),
        "routines": (60, 470), "proxy_runs": (430, 470), "proxy_message_logs": (800, 470), "notices": (1160, 470),
        "leave_requests": (250, 735), "exam_schedules": (640, 735),
    }
    for name, fields in entities.items():
        x, y = positions[name]
        rounded(draw, (x, y, x + 300, y + 170), 16, "#F7FBF8", "#09B884", 3)
        draw.rectangle((x, y, x + 300, y + 38), fill="#04795A")
        draw.text((x + 12, y + 8), name, font=pil_font(20, True), fill="#FFFFFF")
        for index, field in enumerate(fields):
            draw.text((x + 14, y + 48 + index * 22), field, font=pil_font(15), fill="#10211B")
    for a, b in [("institutions", "routines"), ("users", "notices"), ("teacher_profiles", "class_sections"), ("routines", "proxy_runs"), ("proxy_runs", "proxy_message_logs"), ("routines", "leave_requests"), ("routines", "exam_schedules")]:
        ax, ay = positions[a]
        bx, by = positions[b]
        draw.line((ax + 150, ay + 170, bx + 150, by), fill="#5B6B7F", width=3)
    img.save(path)


def draw_results(path):
    img = Image.new("RGB", (1400, 860), "#F7FBF8")
    draw = ImageDraw.Draw(img)
    draw.text((55, 45), "Representative Scholarly Screens", font=pil_font(36, True), fill="#10211B")
    cards = [
        ("Sign in", "Branded auth page with logo, clean card, grid background", 70, 140),
        ("Admin dashboard", "Operational overview for routines, proxy work, notices, teachers, and classes", 510, 140),
        ("Classroom", "Subject feed with assignments, deadlines, tests, files, submissions, and comments", 950, 140),
        ("Proxy manager", "Generate proxy routines, approve day plans, preview WhatsApp messages", 70, 500),
        ("Exam schedule", "Build halls, time slots, subjects, and invigilator assignments", 510, 500),
        ("Mobile sidebar", "Collapsed icon navigation tuned for admin, teacher, and student roles", 950, 500),
    ]
    for name, desc, x, y in cards:
        rounded(draw, (x, y, x + 360, y + 260), 24, "#FFFFFF", "#D9E8DF", 3)
        draw.rectangle((x, y, x + 360, y + 58), fill="#10211B")
        draw.text((x + 24, y + 16), name, font=pil_font(24, True), fill="#8BED9A")
        draw.rounded_rectangle((x + 28, y + 92, x + 332, y + 128), radius=10, fill="#EAF8EF")
        draw.rounded_rectangle((x + 28, y + 148, x + 250, y + 176), radius=8, fill="#BDF8C8")
        draw.rounded_rectangle((x + 28, y + 196, x + 310, y + 218), radius=8, fill="#EEF2F0")
        yy = y + 220
        for line in textwrap.wrap(desc, 34)[:2]:
            draw.text((x + 28, yy), line, font=pil_font(15), fill="#5B6B7F")
            yy += 18
    img.save(path)


def draw_data_flow(path):
    img = Image.new("RGB", (1400, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((45, 35), "Level 0 Data Flow Diagram", font=pil_font(34, True), fill="#10211B")
    rounded(draw, (540, 315, 860, 535), 28, "#F7FBF8", "#09B884", 4)
    center(draw, (540, 315, 860, 535), "Scholarly\nApplication", pil_font(26, True))
    items = [
        ((80, 145, 350, 275), "Admin", "teachers, classes,\nroutines, notices"),
        ((80, 565, 350, 695), "Teacher", "leave, posts,\nassignments"),
        ((1030, 145, 1300, 275), "Student", "classroom views,\nsubmissions"),
        ((1030, 565, 1300, 695), "WhatsApp Cloud API", "routine messages,\nprovider status"),
        ((540, 690, 860, 820), "MySQL + Firebase", "records, snapshots,\nlive classroom posts"),
    ]
    for box, title, body in items:
        rounded(draw, box, 22, "#FFFFFF", "#BDF8C8", 4)
        draw.text((box[0] + 22, box[1] + 20), title, font=pil_font(24, True), fill="#04795A")
        draw.text((box[0] + 22, box[1] + 62), body, font=pil_font(18), fill="#5B6B7F")
    flows = [
        ((350, 210), (540, 370), "manage data"),
        ((350, 630), (540, 475), "submit work"),
        ((1030, 210), (860, 370), "view updates"),
        ((860, 475), (1030, 630), "send messages"),
        ((700, 535), (700, 690), "read/write data"),
    ]
    for start, end, label in flows:
        draw.line([start, end], fill="#04795A", width=5)
        mx = (start[0] + end[0]) // 2
        my = (start[1] + end[1]) // 2
        draw.rounded_rectangle((mx - 85, my - 18, mx + 85, my + 18), radius=12, fill="#FFFFFF", outline="#D9E8DF", width=2)
        center(draw, (mx - 85, my - 18, mx + 85, my + 18), label, pil_font(14))
    img.save(path)


def draw_auth_sequence(path):
    img = Image.new("RGB", (1500, 820), "#F7FBF8")
    draw = ImageDraw.Draw(img)
    draw.text((45, 35), "Authentication Sequence Diagram", font=pil_font(34, True), fill="#10211B")
    actors = [("User", 140), ("Vue Login Page", 430), ("Firebase Auth", 720), ("Laravel Auth", 1010), ("MySQL", 1290)]
    top, bottom = 140, 700
    for name, x in actors:
        rounded(draw, (x - 105, top - 45, x + 105, top + 15), 18, "#FFFFFF", "#09B884", 3)
        center(draw, (x - 105, top - 45, x + 105, top + 15), name, pil_font(18, True))
        draw.line((x, top + 15, x, bottom), fill="#B8CCC1", width=3)
    steps = [
        (140, 430, 210, "enter email and password"),
        (430, 720, 285, "sign in when Firebase account exists"),
        (720, 430, 360, "return ID token"),
        (430, 1010, 435, "submit token or legacy credentials"),
        (1010, 1290, 510, "lookup user and role"),
        (1290, 1010, 585, "return account"),
        (1010, 430, 660, "create Laravel session"),
    ]
    for x1, x2, y, label in steps:
        draw.line((x1, y, x2, y), fill="#04795A", width=4)
        direction = 1 if x2 > x1 else -1
        draw.polygon([(x2, y), (x2 - direction * 16, y - 9), (x2 - direction * 16, y + 9)], fill="#04795A")
        draw.text((min(x1, x2) + 15, y - 28), label, font=pil_font(16), fill="#10211B")
    img.save(path)


def draw_proxy_workflow(path):
    img = Image.new("RGB", (1500, 760), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((45, 35), "Proxy Routine Workflow", font=pil_font(34, True), fill="#10211B")
    steps = [
        ("Select active\nroutine", 70),
        ("Choose absent\nteachers", 325),
        ("Generate proxy\nassignments", 580),
        ("Review and\napprove run", 835),
        ("Preview WhatsApp\nmessages", 1090),
        ("Send and log\nresults", 1345),
    ]
    y = 320
    for index, (label, x) in enumerate(steps):
        rounded(draw, (x - 105, y - 80, x + 105, y + 80), 24, "#F7FBF8", "#09B884", 4)
        center(draw, (x - 105, y - 80, x + 105, y + 80), label, pil_font(19, True))
        draw.ellipse((x - 26, y - 145, x + 26, y - 93), fill="#09B884")
        center(draw, (x - 26, y - 145, x + 26, y - 93), str(index + 1), pil_font(20, True), "#FFFFFF")
        if index < len(steps) - 1:
            nx = steps[index + 1][1]
            draw.line((x + 105, y, nx - 105, y), fill="#04795A", width=5)
            draw.polygon([(nx - 105, y), (nx - 126, y - 12), (nx - 126, y + 12)], fill="#04795A")
    draw.text((85, 565), "The workflow stores a generated proxy grid, teacher schedule snapshots, message previews, send statuses, and provider error messages.", font=pil_font(20), fill="#5B6B7F")
    img.save(path)


def draw_component_map(path):
    img = Image.new("RGB", (1500, 960), "#F7FBF8")
    draw = ImageDraw.Draw(img)
    draw.text((45, 35), "Frontend Component Map", font=pil_font(34, True), fill="#10211B")
    groups = [
        ("App Shell", ["AppLayout.vue", "AppSidebar.vue", "AppTopbar.vue", "ApplicationLogo.vue"], (70, 135)),
        ("Auth Pages", ["Login.vue", "Register.vue", "ForgotPassword.vue", "ResetPassword.vue"], (560, 135)),
        ("Admin Pages", ["Dashboard.vue", "Routines/Create.vue", "Routines/Show.vue", "ProxyManager/Index.vue"], (1050, 135)),
        ("Academic Pages", ["Classroom/Index.vue", "ExamSchedule/Index.vue", "Noticeboard/Index.vue", "LeaveRequests/Index.vue"], (315, 540)),
        ("Shared Logic", ["useDashboardClassroomFeed.js", "Firebase classroom collections", "Inertia forms and router", "Ziggy route helpers"], (805, 540)),
    ]
    for title, files, (x, y) in groups:
        rounded(draw, (x, y, x + 390, y + 300), 24, "#FFFFFF", "#D9E8DF", 3)
        draw.rectangle((x, y, x + 390, y + 52), fill="#10211B")
        draw.text((x + 22, y + 14), title, font=pil_font(22, True), fill="#8BED9A")
        yy = y + 82
        for f in files:
            draw.rounded_rectangle((x + 24, yy, x + 366, yy + 38), radius=12, fill="#EAF8EF")
            draw.text((x + 40, yy + 9), f, font=pil_font(17), fill="#10211B")
            yy += 52
    for start, end in [((460, 285), (560, 285)), ((950, 285), (1050, 285)), ((265, 435), (510, 540)), ((1245, 435), (1000, 540)), ((705, 435), (1000, 540))]:
        draw.line([start, end], fill="#04795A", width=4)
    img.save(path)


def draw_class_model(path):
    img = Image.new("RGB", (1500, 980), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((45, 35), "Simplified Model Relationship Diagram", font=pil_font(34, True), fill="#10211B")
    models = [
        ("Institution", ["has many users", "has many teachers", "has many classes"], (80, 135)),
        ("User", ["belongs to institution", "role: admin/teacher/student", "has settings"], (565, 135)),
        ("TeacherProfile", ["belongs to user", "has WhatsApp number", "appears in routines"], (1050, 135)),
        ("Routine", ["belongs to institution", "stores generated grid", "has proxy runs"], (80, 540)),
        ("ProxyRun", ["belongs to routine", "stores assignments", "has message logs"], (565, 540)),
        ("ClassSection", ["belongs to institution", "has subjects JSON", "has join code"], (1050, 540)),
    ]
    for name, fields, (x, y) in models:
        rounded(draw, (x, y, x + 360, y + 255), 18, "#F7FBF8", "#09B884", 3)
        draw.rectangle((x, y, x + 360, y + 48), fill="#04795A")
        draw.text((x + 20, y + 13), name, font=pil_font(22, True), fill="#FFFFFF")
        yy = y + 78
        for field in fields:
            draw.text((x + 26, yy), "+ " + field, font=pil_font(18), fill="#10211B")
            yy += 44
    for start, end in [((440, 260), (565, 260)), ((925, 260), (1050, 260)), ((260, 390), (260, 540)), ((745, 390), (745, 540)), ((1230, 390), (1230, 540)), ((440, 665), (565, 665)), ((1050, 665), (925, 665))]:
        draw.line([start, end], fill="#5B6B7F", width=4)
    img.save(path)


def add_long_content(doc):
    h(doc, "Acknowledgement", 1)
    p(doc, "This project was completed as a practical full-stack web application for managing institutional academic operations. The development process involved database design, role-based dashboards, real-time classroom communication, routine generation, proxy management, and external service integration.")
    p(doc, "The application was developed with attention to real school workflows rather than a narrow CRUD-only pattern. This documentation records the design decisions, implementation details, testing results, current limitations, and future improvements needed to move the system from a local prototype into a production-ready product.")
    doc.add_page_break()

    h(doc, "Table of Contents", 1)
    for item in [
        "Chapter 1: Introduction",
        "Chapter 2: System Analysis and Requirements",
        "Chapter 3: System Design",
        "Chapter 4: Technology Stack and Environment Setup",
        "Chapter 5: Implementation",
        "Chapter 6: Testing",
        "Chapter 7: Results and Screenshots",
        "Chapter 8: Limitations and Future Work",
        "Chapter 9: Conclusion",
        "Chapter 10: References and Appendices",
        "Appendix A: Complete Design Artifacts",
        "Appendix B: Detailed Route and Data Dictionary",
        "Appendix C: Test Evidence and Submission Checklist",
    ]:
        p(doc, item, after=3)
    doc.add_page_break()

    h(doc, "Executive Summary", 1)
    p(doc, "Scholarly is a web-based school operations platform designed to help an institution manage academic work that can become fragmented across paper registers, spreadsheet files, chat groups, and disconnected manual processes. The system focuses on routine generation, teacher and classroom management, proxy class planning, leave requests, noticeboard communication, exam scheduling, classroom assignments, class tests, and WhatsApp routine update messaging.")
    p(doc, "The project uses Laravel 12 as the backend framework, Vue 3 as the frontend framework, Inertia.js as the bridge between server-side routing and client-side screens, MySQL for persistent institutional data, Firebase for live classroom streams, and Vite for the build pipeline. The current local dataset includes 20 teacher profiles, 10 class sections, 1 active routine, 8 noticeboard posts, 3 leave requests, 1 exam schedule, 2 proxy runs, and 12 seeded classroom posts in Firebase.")

    h(doc, "Chapter 1: Introduction", 1)
    h(doc, "1.1 Background", 2)
    p(doc, "Small and mid-sized schools often manage routine creation, proxy class assignment, notices, exam schedules, classroom posts, and teacher leave through separate manual processes. A routine may be stored in a spreadsheet, notices may be shared verbally or through messaging groups, and proxy work may be rearranged at the last moment with limited visibility.")
    p(doc, "Scholarly provides a single web application where an institution can define teachers and class sections, create an academic routine, generate proxy plans for absent teachers, publish notices, schedule exams, manage leave requests, and operate classroom feeds for assignments and class tests.")
    h(doc, "1.2 Problem Statement", 2)
    p(doc, "Many school workflows are handled across disconnected tools. Because the data is not connected, changes in one area do not automatically reflect in the rest of the system. Scholarly solves this by centralizing academic operations in a structured Laravel and Vue application with role-based access and live classroom updates.")
    h(doc, "1.3 Objectives", 2)
    p(doc, "Main objective: build a web-based institution management system that allows school administrators, teachers, and students to coordinate routines, classrooms, notices, leave, exams, and proxy work from a unified platform.")
    bullets(doc, [
        "Implement secure registration and login with Laravel authentication and Firebase-aware account flows.",
        "Provide role-based dashboards for administrators, teachers, and students.",
        "Allow administrators to manage teachers, class sections, join codes, and institution settings.",
        "Create, import, activate, rename, regenerate, and delete academic routines.",
        "Support proxy routine planning for absent teachers, including approval and WhatsApp update preparation.",
        "Provide noticeboard and staff notice workflows with urgency, visibility, acknowledgement, and read-count tracking.",
        "Support teacher leave requests, exam scheduling, live classroom feeds, and mobile-friendly navigation.",
    ])
    h(doc, "1.4 Scope and Limitations", 2)
    p(doc, "The current scope includes institution setup, users, teacher profiles, class sections, routine generation, proxy planning, exam scheduling, leave management, noticeboard communication, classroom activity, dashboard summaries, search, settings, and WhatsApp Cloud API message logging. The system does not yet include payroll, biometric attendance, production deployment automation, or complete report export for every module.")
    h(doc, "1.5 Document Organization", 2)
    p(doc, "The remaining chapters describe requirements, design, technology stack, implementation, testing, results, limitations, conclusion, references, and appendices.")

    h(doc, "Chapter 2: System Analysis and Requirements", 1)
    h(doc, "2.1 Existing System Study", 2)
    p(doc, "The closest existing process is a manual school administration workflow. Routines are commonly prepared using spreadsheets or paper tables. Teacher information is maintained separately from routine data. Proxy classes are often arranged through direct messages or informal staff communication. This creates duplicated work and weak visibility.")
    h(doc, "2.2 Requirement Gathering Method", 2)
    p(doc, "Requirements were gathered through iterative project development, direct testing of school administration scenarios, and repeated interface feedback. The project was adjusted after practical issues appeared, such as mobile sidebar spacing, teacher-name synchronization, login behavior after Firebase was added, and WhatsApp Cloud API error handling.")
    h(doc, "2.3 Functional Requirements", 2)
    table(doc, [
        ("ID", "Requirement", "User Role"),
        ("FR-01", "Register an administrator, teacher, or student account.", "Guest"),
        ("FR-02", "Log in and route the user to the correct role dashboard.", "All users"),
        ("FR-03", "Manage institution profile, teachers, class sections, and settings.", "Admin"),
        ("FR-04", "Create, import, activate, rename, regenerate, and delete routines.", "Admin"),
        ("FR-05", "Generate proxy plans and prepare WhatsApp routine messages.", "Admin"),
        ("FR-06", "Publish institutional and staff notices.", "Admin, Teacher"),
        ("FR-07", "Submit and review leave requests.", "Teacher, Admin"),
        ("FR-08", "Create exam schedules with halls, time slots, subjects, and invigilators.", "Admin"),
        ("FR-09", "View classroom feeds, assignments, class tests, comments, and submissions.", "Teacher, Student"),
        ("FR-10", "Use responsive role-aware navigation on desktop and mobile.", "All users"),
    ], (0.8, 4.2, 1.5))
    h(doc, "2.4 Non-Functional Requirements", 2)
    kv(doc, [
        ("Performance", "Pages should load quickly on local development and normal broadband connections. Vite builds production assets for optimized delivery."),
        ("Security", "Passwords are hashed by Laravel. Authenticated routes are protected by middleware. Firebase tokens are used where live services are involved."),
        ("Usability", "The interface uses a consistent sidebar, role-specific menus, clear forms, responsive cards, and mobile adjustments."),
        ("Maintainability", "The backend uses Laravel controllers, services, Eloquent models, migrations, and configuration files. The frontend uses Vue single-file components and composables."),
        ("Reliability", "External WhatsApp failures are caught and logged instead of crashing the application."),
    ])
    h(doc, "2.5 Use Case Diagram", 2)
    doc.add_picture(str(ASSETS / "usecase.png"), width=Inches(6.4))
    caption(doc, "Figure 1: Main use cases for administrators, teachers, and students.")
    figure_description(doc, "This figure identifies the three primary user roles and the system actions available to each role. It shows that Scholarly is centered on administrative setup, teacher classroom work, student participation, and shared schedule updates.")
    table(doc, [
        ("Actor", "Use Cases", "Expected Outcome"),
        ("Admin", "Manage institution, teachers, class sections, routines, proxy plans, notices, exam schedules, and settings.", "The institution can run the academic workflow from one dashboard."),
        ("Teacher", "View assigned schedule, post classroom work, create assignments and class tests, request leave, read notices, and receive routine updates.", "Teachers can manage class work and remain updated about schedule changes."),
        ("Student", "Join class section, view classroom posts, check deadlines, submit work, read notices, and view schedules.", "Students can follow academic work without relying on separate chat groups."),
    ], (1.15, 3.9, 1.75))
    h(doc, "2.6 Data Flow Diagram", 2)
    doc.add_picture(str(ASSETS / "data-flow.png"), width=Inches(6.4))
    caption(doc, "Figure 2: Level 0 data flow between users, Scholarly, storage, and WhatsApp Cloud API.")
    figure_description(doc, "This figure explains how information enters and leaves the application. Administrators, teachers, and students interact with Scholarly, while the system reads and writes records to MySQL and Firebase and sends routine messages through the WhatsApp Cloud API.")

    h(doc, "Chapter 3: System Design", 1)
    h(doc, "3.1 System Architecture", 2)
    p(doc, "Scholarly uses a Laravel monolith with an Inertia-powered Vue frontend. Laravel owns routing, authorization, validation, database persistence, and server-side responses. Vue handles interactive pages in resources/js and communicates through Inertia visits, Axios requests, and Firebase listeners where real-time classroom data is required.")
    doc.add_picture(str(ASSETS / "architecture.png"), width=Inches(6.4))
    caption(doc, "Figure 3: High-level architecture of Scholarly.")
    figure_description(doc, "This figure shows the layered application structure. Vue and Inertia form the user interface, Laravel handles controllers and services, MySQL and Firebase store data, and the WhatsApp integration sends approved routine updates.")
    h(doc, "3.2 Entity Relationship Diagram", 2)
    doc.add_picture(str(ASSETS / "erd.png"), width=Inches(6.4))
    caption(doc, "Figure 4: Core relational entities used by the Laravel backend.")
    figure_description(doc, "This figure maps the main database entities and their relationships. It highlights how institutions connect to users, teacher profiles, class sections, routines, notices, leave requests, exam schedules, proxy runs, and WhatsApp message logs.")
    h(doc, "3.3 Model Relationship Diagram", 2)
    doc.add_picture(str(ASSETS / "model-relationships.png"), width=Inches(6.4))
    caption(doc, "Figure 5: Simplified model relationships used by the application layer.")
    figure_description(doc, "This figure presents the same domain in a simpler model-oriented view. It focuses on the Laravel application relationships that developers use when retrieving institution, user, teacher, routine, proxy, and class-section data.")
    h(doc, "3.4 Database Schema", 2)
    table(doc, [
        ("Table", "Main Columns", "Purpose"),
        ("users", "name, email, firebase_uid, password, role, institution_id, settings", "Stores login accounts and role links."),
        ("institutions", "owner_user_id, name, short_name, phone, email, address, academic_year, settings", "Stores institution-level profile and defaults."),
        ("teacher_profiles", "institution_id, user_id, name, whatsapp_number, join_code, status", "Stores teacher directory and WhatsApp details."),
        ("class_sections", "institution_id, class_teacher_profile_id, class_name, section_name, join_code, subjects", "Stores class-section directory and assigned subjects."),
        ("routines", "institution_id, name, status, days, periods, classes, teachers, generated_grid", "Stores generated academic routines and schedule snapshots."),
        ("proxy_runs", "routine_id, date, day_label, status, absent_teachers, assignments, proxy_generated_grid", "Stores proxy routine outputs."),
        ("notices", "institution_id, user_id, board, title, message, urgency, visibility", "Stores institutional and staff notices."),
        ("leave_requests", "institution_id, routine_id, teacher_id, start_date, end_date, status", "Stores leave request workflow data."),
        ("exam_schedules", "institution_id, routine_id, halls, time_slots, exam_grid, status", "Stores exam scheduling data."),
        ("proxy_message_logs", "proxy_run_id, teacher_profile_id, whatsapp_number, status, provider_message_id", "Stores WhatsApp send history."),
    ], (1.25, 3.15, 2.1))
    h(doc, "3.5 Route and API Design", 2)
    p(doc, "The project primarily uses web routes with Inertia responses instead of a separate public REST API. Selected JSON endpoints support notifications, search, and WhatsApp previews. The active route list contains 79 routes.")
    table(doc, [
        ("Method", "Endpoint", "Purpose", "Auth"),
        ("GET", "/dashboard", "Render role-specific dashboard.", "Yes"),
        ("GET, POST", "/login", "Render login screen and create session.", "No"),
        ("GET, POST", "/register", "Render registration screen and create account.", "No"),
        ("GET, POST", "/routines", "List and create routines.", "Yes"),
        ("GET, POST", "/proxy-manager", "List and generate proxy runs.", "Yes"),
        ("GET", "/proxy-manager/{proxyRun}/whatsapp-preview", "Return prepared WhatsApp messages as JSON.", "Yes"),
        ("POST", "/proxy-manager/{proxyRun}/whatsapp", "Send WhatsApp routine updates and log results.", "Yes"),
        ("GET, POST", "/noticeboard", "List and publish notices.", "Yes"),
        ("GET, POST", "/leave-requests", "List and create leave requests.", "Yes"),
        ("GET, POST", "/exam-schedule", "List and create exam schedules.", "Yes"),
        ("GET", "/classroom", "Render classroom stream context.", "Yes"),
        ("GET", "/notifications", "Return notification center JSON payload.", "Yes"),
        ("GET", "/search/features", "Return searchable feature and classroom items.", "Yes"),
    ], (0.85, 2.05, 2.8, 0.8))
    h(doc, "3.6 Frontend Component Structure", 2)
    code(doc, """resources/js
  app.js
  Layouts/AppLayout.vue
  Layouts/GuestLayout.vue
  Components/AppSidebar.vue
  Components/AppTopbar.vue
  Components/ApplicationLogo.vue
  Pages/Auth/Login.vue
  Pages/Auth/Register.vue
  Pages/Dashboard.vue
  Pages/TeacherDashboard.vue
  Pages/StudentDashboard.vue
  Pages/Routines/Create.vue
  Pages/Routines/Show.vue
  Pages/ProxyManager/Index.vue
  Pages/Classroom/Index.vue
  Pages/Noticeboard/Index.vue
  Pages/ExamSchedule/Index.vue
  Pages/Settings/Index.vue
  Composables/useDashboardClassroomFeed.js""")
    doc.add_picture(str(ASSETS / "component-map.png"), width=Inches(6.4))
    caption(doc, "Figure 6: Frontend component grouping for the Vue and Inertia interface.")
    figure_description(doc, "This figure organizes the main Vue files by responsibility. It separates the app shell, authentication screens, admin pages, academic pages, and shared frontend logic so the interface structure is easier to understand.")
    h(doc, "3.7 Sequence Diagram: Authentication", 2)
    doc.add_picture(str(ASSETS / "auth-sequence.png"), width=Inches(6.4))
    caption(doc, "Figure 7: Authentication flow for Firebase-aware and legacy local accounts.")
    figure_description(doc, "This figure traces the login process from the user interface to Firebase, Laravel, and MySQL. It shows how Scholarly supports both Firebase-backed accounts and older local accounts that existed before Firebase was added.")
    h(doc, "3.8 Workflow Diagram: Proxy Routine", 2)
    doc.add_picture(str(ASSETS / "proxy-workflow.png"), width=Inches(6.4))
    caption(doc, "Figure 8: Proxy routine generation, approval, preview, and WhatsApp logging flow.")
    figure_description(doc, "This figure explains the proxy-management workflow from routine selection to final message logging. It shows how an administrator selects absent teachers, generates assignments, approves the run, previews WhatsApp messages, and records delivery results.")
    h(doc, "3.9 Wireframes and UI Mockups", 2)
    doc.add_picture(str(ASSETS / "results.png"), width=Inches(6.4))
    caption(doc, "Figure 9: Branded representation of the main screens included in the final application.")
    figure_description(doc, "This figure gives a compact visual overview of the major Scholarly screens. It represents the sign-in page, admin dashboard, classroom feed, proxy manager, exam schedule, and responsive mobile sidebar.")

    h(doc, "Chapter 4: Technology Stack and Environment Setup", 1)
    h(doc, "4.1 Technology Stack", 2)
    table(doc, [
        ("Layer", "Technology", "Version or Package", "Purpose"),
        ("Backend framework", "Laravel", "12.x", "Routing, controllers, validation, services, Eloquent ORM, sessions."),
        ("Language", "PHP", "8.2.12", "Backend runtime."),
        ("Frontend framework", "Vue.js", "3.4.x", "Interactive single-file components."),
        ("Bridge", "Inertia.js", "2.x", "Connects Laravel routes with Vue pages."),
        ("Database", "MySQL", "Local MySQL through XAMPP", "Persistent relational storage."),
        ("Authentication", "Laravel auth, Sanctum, Firebase", "Sanctum 4.x, Firebase 12.x", "Session auth plus Firebase-aware flows."),
        ("Styling", "Tailwind CSS", "3.2.x", "Responsive utility-first UI styling."),
        ("Build tool", "Vite", "7.x", "Development server and production asset compilation."),
        ("External service", "WhatsApp Cloud API", "Graph API v25.0", "Routine update messaging."),
    ], (1.35, 1.35, 1.7, 2.1))
    h(doc, "4.2 Development Environment", 2)
    kv(doc, [
        ("Operating system", "Windows development environment"),
        ("PHP", "8.2.12"),
        ("Composer", "2.8.12"),
        ("Node.js", "24.16.0"),
        ("npm", "11.13.0"),
        ("Local server", "php artisan serve at http://127.0.0.1:8000"),
        ("Frontend server", "Vite on port 5174 when hot reload is enabled"),
        ("Database", "MySQL, configured through Laravel .env"),
    ])
    h(doc, "4.3 Installation Steps", 2)
    code(doc, """git clone <repository-url> scholarly
cd scholarly
composer install
cmd /c npm install
copy .env.example .env
php artisan key:generate
php artisan migrate --seed
php artisan serve
cmd /c npm run dev""")
    h(doc, "4.4 Environment Variables", 2)
    code(doc, """APP_NAME=Scholarly
APP_ENV=local
APP_KEY=<generated application key>
APP_URL=http://127.0.0.1:8000
DB_CONNECTION=mysql
DB_DATABASE=<database name>
FIREBASE_API_KEY=<firebase web api key>
FIREBASE_AUTH_DOMAIN=<firebase auth domain>
FIREBASE_PROJECT_ID=<firebase project id>
WHATSAPP_API_VERSION=v25.0
WHATSAPP_PHONE_NUMBER_ID=<whatsapp phone number id>
WHATSAPP_ACCESS_TOKEN=<token hidden>
WHATSAPP_DRY_RUN=false""")

    h(doc, "Chapter 5: Implementation", 1)
    h(doc, "5.1 Backend Implementation", 2)
    p(doc, "The backend is organized around Laravel controllers and service classes. Controllers validate user input, enforce route access through middleware, and return Inertia pages or redirects. Eloquent models represent persistent entities and define casts for JSON fields such as routine grids, teacher schedules, proxy assignments, institution settings, and notification settings.")
    h(doc, "Authentication and role routing", 3)
    p(doc, "Authentication is based on Laravel Breeze-style session login with additional Firebase-aware behavior. The login page can submit a Firebase ID token for Firebase-backed accounts. Local legacy users can still authenticate through the Laravel database, which was necessary because the institute dataset existed before Firebase was added.")
    code(doc, """// resources/js/app.js
createInertiaApp({
    resolve: (name) => resolvePageComponent(
        `./Pages/${name}.vue`,
        import.meta.glob('./Pages/**/*.vue')
    ),
    setup({ el, App, props, plugin }) {
        return createApp({ render: () => h(App, props) })
            .use(plugin)
            .use(ZiggyVue)
            .mount(el);
    },
});""")
    h(doc, "Routine and proxy management", 3)
    p(doc, "The routine module stores school days, periods, classes, teachers, generation rules, generated grid data, teacher schedules, and metrics. ProxyRunController uses a ProxyEngine service to create a substitute plan for absent teachers, then stores both the proxy grid and teacher schedule snapshots.")
    code(doc, """// ProxyRunController@store, simplified
$routine = Routine::findOrFail($data['routineId']);
$generated = $engine->generate($routine, $data);
$proxyGrid = $this->proxyGridForRun($routine, $generated['day'], $generated['assignments']);
$proxySchedule = $this->teacherScheduleFromGrid($routine, $proxyGrid);

ProxyRun::create([
    'routine_id' => $routine->id,
    'day_label' => $generated['day'],
    'assignments' => $generated['assignments'],
    'proxy_generated_grid' => $proxyGrid,
    'proxy_teacher_schedule' => $proxySchedule,
    'metrics' => $generated['metrics'],
]);""")
    h(doc, "WhatsApp routine messaging", 3)
    p(doc, "WhatsAppRoutineMessenger prepares teacher-specific schedule messages from an approved proxy run. It normalizes Bangladeshi phone numbers, skips teachers without numbers, respects personal WhatsApp notification settings, catches connection errors, records provider IDs, and stores failure messages returned by Meta.")
    h(doc, "5.2 Frontend Implementation", 2)
    p(doc, "The frontend is built with Vue single-file components under resources/js. AppLayout.vue and AppSidebar.vue provide the shell for authenticated pages. GuestLayout.vue wraps login, registration, password reset, and verification pages. Each page component receives props from Laravel through Inertia and submits changes with Inertia router calls or forms.")
    code(doc, """// ProxyManager/Index.vue, simplified
router.post(`/proxy-manager/${run.id}/whatsapp`, {
    messages: whatsappDialog.value.messages,
}, {
    preserveScroll: true,
    onSuccess: () => closeWhatsAppDialog(),
    onError: () => {
        whatsappDialog.value.error = 'Could not send WhatsApp messages.';
    },
});""")
    h(doc, "5.3 Authentication Flow", 2)
    bullets(doc, [
        "The user opens /login and enters email and password.",
        "The Vue login component checks Firebase configuration and legacy account status.",
        "For Firebase accounts, the component signs in through Firebase and sends the ID token to Laravel.",
        "For legacy local accounts, the component submits credentials directly to Laravel.",
        "Laravel validates the account, creates a session, and redirects to /dashboard.",
        "DashboardController detects the role and returns the correct dashboard data.",
    ], numbered=True)
    h(doc, "5.4 Notable Challenges and Solutions", 2)
    table(doc, [
        ("Challenge", "Resolution"),
        ("Teacher name synchronization", "The update flow was expanded to synchronize related JSON snapshots in routines, proxy runs, leave records, and exam schedules."),
        ("Collapsed sidebar spacing", "The sidebar spacing was adjusted by role and collapsed state so icons fit without requiring a scrollable collapsed sidebar."),
        ("Firebase migration and legacy login", "A local-only legacy login path was added for users with no firebase_uid."),
        ("WhatsApp API failures", "The messenger service now catches timeouts, OAuth errors, and provider failures, then logs the response."),
        ("Vite hot file blank page", "Removing stale public/hot markers forces Laravel back to built assets when Vite is not needed."),
    ], (1.75, 4.75))

    h(doc, "Chapter 6: Testing", 1)
    h(doc, "6.1 Testing Approach", 2)
    p(doc, "Testing used manual click-through testing, Laravel feature tests, PHP syntax checks, route inspection, Vite production builds, and targeted database checks. The latest production frontend build completed successfully. The Laravel test suite reported 43 passing tests and 3 failing tests.")
    h(doc, "6.2 Test Cases", 2)
    table(doc, [
        ("Test ID", "Description", "Input or Action", "Expected Result", "Status"),
        ("TC-01", "Login page rendering", "GET /login", "Login screen loads.", "Pass"),
        ("TC-02", "Admin dashboard access", "Authenticated admin opens /dashboard", "Operations dashboard is returned.", "Pass"),
        ("TC-03", "Teacher dashboard context", "Authenticated teacher opens /dashboard", "Teacher schedule and classroom context are returned.", "Pass"),
        ("TC-04", "Student dashboard isolation", "Authenticated student opens /dashboard", "Only the student class context is returned.", "Pass"),
        ("TC-05", "Search permissions", "Teacher and student search classrooms", "Only allowed classroom subjects are shown.", "Pass"),
        ("TC-06", "Settings permissions", "Teacher or student attempts admin cleanup", "Request is blocked.", "Pass"),
        ("TC-07", "WhatsApp opt-out", "Teacher disables WhatsApp updates", "Routine message is skipped.", "Pass"),
        ("TC-08", "Production frontend build", "cmd /c npm run build", "Vite compiles assets successfully.", "Pass"),
        ("TC-09", "Default auth login test", "PHPUnit Breeze login test", "Expected authenticated state.", "Fail - update for Firebase-aware flow"),
        ("TC-10", "Default registration test", "PHPUnit Breeze registration test", "Expected authenticated state.", "Fail - update for Firebase token requirement"),
        ("TC-11", "Default root response test", "GET /", "Expected 200.", "Fail - app redirects by design"),
    ], (0.7, 1.65, 1.45, 1.75, 0.95))
    h(doc, "6.3 Bug Tracking", 2)
    p(doc, "Significant fixes included teacher-name propagation across schedule snapshots, WhatsApp timeout and request exception handling, phone-number normalization for Bangladesh numbers, stale Vite hot-file recovery, and mobile sidebar spacing adjustments.")

    h(doc, "Chapter 7: Results and Screenshots", 1)
    p(doc, "The current prototype demonstrates a functional multi-role school operations system with realistic data. The local institute dataset contains 14 users, 20 teachers, 10 class sections, 1 active academic routine, 8 notices, 3 leave requests, 1 exam schedule, 2 proxy runs, 227 WhatsApp message log records, and 12 Firebase classroom posts seeded for presentation.")
    doc.add_picture(str(ASSETS / "results.png"), width=Inches(6.4))
    caption(doc, "Figure 10: Representative screen set for the final Scholarly interface.")
    figure_description(doc, "This figure summarizes the visible result of the implementation. It is included in the results chapter to connect the designed modules with the screens that users interact with during normal use.")
    table(doc, [
        ("Screen", "What the user sees", "Notable result"),
        ("Login", "Large Scholarly brand area, fading grid background, white sign-in card.", "Local legacy account and Firebase-aware account paths are supported."),
        ("Registration", "Role selector for admin, teacher, and student, plus institution or join-code fields.", "The form supports different onboarding paths by role."),
        ("Admin dashboard", "Summary cards, routine status, proxy work, notices, analytics links.", "Administrators can reach all operational modules from one shell."),
        ("Routine view", "Generated class grid and teacher schedule.", "Routine data is stored as structured JSON snapshots."),
        ("Proxy manager", "Absent teacher selection, generated proxy assignments, approval state, WhatsApp preview.", "Approved runs can generate teacher-specific routine messages."),
        ("Classroom", "Subject feed with updates, assignments, class tests, comments, and submissions.", "Live posts use Firebase collections scoped by institution."),
        ("Noticeboard", "Institutional and staff notices with urgency and acknowledgement.", "Visibility rules separate all-user and teacher-only communication."),
        ("Exam schedule", "Halls, time slots, subjects, invigilator options, and generated exam grid.", "Exam schedule data is reusable and editable."),
    ], (1.35, 2.7, 2.45))

    h(doc, "Chapter 8: Limitations and Future Work", 1)
    h(doc, "8.1 Current Limitations", 2)
    bullets(doc, [
        "WhatsApp routine messages currently depend on Meta allowed recipient rules during development mode.",
        "Some default authentication tests need to be rewritten to reflect the current Firebase-aware authentication flow.",
        "Classroom posts are stored in Firebase, so offline classroom access is limited without a local synchronization layer.",
        "The application is currently configured for local development rather than a hardened production deployment.",
        "Automated browser screenshot capture is not yet part of the test pipeline.",
        "Advanced exports, printed reports, and PDF generation are not implemented for every module.",
    ])
    h(doc, "8.2 Future Work", 2)
    bullets(doc, [
        "Implement WhatsApp utility templates for routine updates and proxy notifications.",
        "Add production deployment configuration with queue workers, scheduled jobs, environment separation, and monitoring.",
        "Add full automated browser tests for login, routine generation, proxy approval, notice creation, and classroom posting.",
        "Add report exports for routines, proxy plans, leave summaries, exam schedules, and classroom deadlines.",
        "Add parent accounts, attendance, marks entry, result publication, and guardian-facing notifications.",
        "Improve classroom offline resilience by mirroring essential classroom post metadata to MySQL.",
        "Add multi-language support for Bangla and English interfaces.",
    ])

    h(doc, "Chapter 9: Conclusion", 1)
    p(doc, "Scholarly was built as a practical full-stack school operations platform, not only as a demonstration of CRUD pages. The system connects institution setup, users, teacher profiles, class sections, routines, proxy planning, notices, leave requests, exam schedules, classroom activity, settings, notifications, and WhatsApp message logging into one role-based application.")
    p(doc, "The project met its main objective by delivering a working Laravel and Vue application that can support administrative, teacher, and student workflows. The implementation also exposed realistic engineering concerns: data migration after Firebase was introduced, stale generated schedule snapshots after teacher changes, external API failure handling, responsive sidebar behavior, and development-mode WhatsApp limitations.")
    p(doc, "The most valuable lesson from the project is that institutional systems depend on consistency across modules. A teacher name, routine cell, notice visibility rule, or proxy message can affect several screens at once. Building Scholarly required attention to that connected behavior.")

    h(doc, "Chapter 10: References and Appendices", 1)
    h(doc, "10.1 References", 2)
    bullets(doc, [
        "Laravel Documentation. https://laravel.com/docs",
        "Vue.js Documentation. https://vuejs.org/guide",
        "Inertia.js Documentation. https://inertiajs.com",
        "Tailwind CSS Documentation. https://tailwindcss.com/docs",
        "Firebase Documentation. https://firebase.google.com/docs",
        "WhatsApp Business Platform Documentation. https://developers.facebook.com/docs/whatsapp",
        "Vite Documentation. https://vite.dev/guide",
        "Lucide Icons Documentation. https://lucide.dev",
    ])
    h(doc, "10.2 Appendices", 2)
    h(doc, "Appendix A: Project Metrics", 3)
    kv(doc, [
        ("Users", "14"), ("Teacher profiles", "20"), ("Class sections", "10"),
        ("Routines", "1"), ("Active routine", "New Academic Routine"),
        ("Active routine days", "5"), ("Active routine periods", "9"),
        ("Notices", "8"), ("Leave requests", "3"), ("Exam schedules", "1"),
        ("Proxy runs", "2"), ("WhatsApp logs", "227"),
        ("Firebase classroom posts", "12 seeded demo posts"),
    ])
    h(doc, "Appendix B: Build and Test Summary", 3)
    kv(doc, [
        ("Frontend build", "cmd /c npm run build completed successfully."),
        ("Laravel tests", "43 passed, 3 failed."),
        ("Known failing tests", "Default auth login, default registration, and default root response expectations need updates."),
        ("Route inventory", "php artisan route:list reported 79 routes."),
    ])
    h(doc, "Appendix C: Important Commands", 3)
    code(doc, """php artisan route:list
php artisan test
php artisan config:clear
php artisan view:clear
cmd /c npm run build
php artisan serve""")
    h(doc, "Final Submission Checklist", 1)
    bullets(doc, [
        "Replace Student ID and Supervisor on the cover page before final submission.",
        "Open the DOCX in Microsoft Word and update the table of contents if page numbers are required by the institution.",
        "Add real browser screenshots if the evaluator requires literal screenshots rather than screen descriptions.",
        "Do not expose real Firebase or WhatsApp secrets in the submitted document.",
        "Export to PDF after final edits and confirm formatting before submission.",
    ])

    h(doc, "Appendix A: Complete Design Artifacts", 1)
    p(doc, "This appendix repeats the design artifacts as standalone deliverables so the documentation remains self-contained even when reviewed separately from the source code.")
    for title, image, label, description in [
        ("A.1 Use Case Diagram", "usecase.png", "Appendix Figure A1: Complete use case diagram.", "The complete use case diagram shows the user roles and the main actions each role performs inside Scholarly."),
        ("A.2 Level 0 Data Flow Diagram", "data-flow.png", "Appendix Figure A2: Complete data flow diagram.", "The complete data flow diagram shows how user actions, database storage, Firebase classroom records, and WhatsApp delivery interact."),
        ("A.3 System Architecture Diagram", "architecture.png", "Appendix Figure A3: Complete architecture diagram.", "The complete architecture diagram shows the frontend, backend, storage, domain modules, authentication layer, and external integration layer."),
        ("A.4 Entity Relationship Diagram", "erd.png", "Appendix Figure A4: Complete entity relationship diagram.", "The complete ERD shows the core database tables and the relationships that support Scholarly's institution, routine, classroom, and messaging modules."),
        ("A.5 Model Relationship Diagram", "model-relationships.png", "Appendix Figure A5: Complete model relationship diagram.", "The complete model relationship diagram simplifies the database view into the main Laravel model relationships used in implementation."),
        ("A.6 Authentication Sequence Diagram", "auth-sequence.png", "Appendix Figure A6: Complete authentication sequence diagram.", "The complete sequence diagram explains how login requests move between the user, Vue, Firebase, Laravel, and MySQL."),
        ("A.7 Proxy Routine Workflow Diagram", "proxy-workflow.png", "Appendix Figure A7: Complete proxy routine workflow diagram.", "The complete workflow diagram shows the full proxy routine process from choosing absences to sending and logging WhatsApp messages."),
        ("A.8 Frontend Component Map", "component-map.png", "Appendix Figure A8: Complete component map.", "The complete component map groups the Vue files by responsibility so the frontend structure can be reviewed quickly."),
        ("A.9 UI Mockup Board", "results.png", "Appendix Figure A9: Complete UI mockup board.", "The complete UI mockup board summarizes the major screens and demonstrates the visual direction of the Scholarly interface."),
    ]:
        doc.add_page_break()
        h(doc, title, 2)
        doc.add_picture(str(ASSETS / image), width=Inches(6.5))
        caption(doc, label)
        figure_description(doc, description)

    doc.add_page_break()
    h(doc, "Appendix B: Detailed Route and Data Dictionary", 1)
    h(doc, "B.1 Detailed Data Dictionary", 2)
    table(doc, [
        ("Entity", "Attribute", "Type", "Description"),
        ("users", "firebase_uid", "string, nullable", "Links a Laravel user to a Firebase authentication account."),
        ("users", "role", "string", "Controls admin, teacher, or student navigation and permissions."),
        ("users", "settings", "json", "Stores notification preferences and user-level UI settings."),
        ("institutions", "academic_year", "string", "Current academic year label shown in institution-level workflows."),
        ("institutions", "settings", "json", "Stores institute-wide preferences, branding, and notification defaults."),
        ("teacher_profiles", "whatsapp_number", "string, nullable", "Bangladesh-format teacher phone number used for WhatsApp routine updates."),
        ("class_sections", "subjects", "json", "List of subjects connected to a class section."),
        ("routines", "generated_grid", "json", "Generated weekly class routine grid."),
        ("routines", "teacher_schedule", "json", "Teacher-specific schedule snapshot derived from the routine grid."),
        ("proxy_runs", "assignments", "json", "Generated proxy class assignments for an absence event."),
        ("proxy_message_logs", "status", "string", "Message status such as sent, skipped, or failed."),
        ("notices", "visibility", "string", "Controls whether notices target all users, teachers, or selected audiences."),
        ("exam_schedules", "exam_grid", "json", "Final exam schedule grid including subjects, halls, and invigilators."),
    ], (1.25, 1.45, 1.35, 2.7))
    h(doc, "B.2 Route Coverage Matrix", 2)
    table(doc, [
        ("Module", "Routes", "Controller or Page"),
        ("Authentication", "/login, /register, /forgot-password, /reset-password", "Auth controllers and Vue auth pages"),
        ("Dashboard", "/dashboard", "DashboardController with role-specific Vue pages"),
        ("Routines", "/routines, /routines/create, /routines/{routine}", "RoutineController and Routines Vue pages"),
        ("Proxy Manager", "/proxy-manager, /proxy-manager/{proxyRun}/whatsapp-preview, /proxy-manager/{proxyRun}/whatsapp", "ProxyRunController and WhatsAppRoutineMessenger"),
        ("Teachers", "/teachers", "Teacher directory management"),
        ("Class Sections", "/classrooms", "Class section directory management"),
        ("Classroom Feed", "/classroom", "Firebase-backed classroom stream"),
        ("Noticeboard", "/noticeboard, /staffroom", "Notice creation, visibility, and acknowledgement"),
        ("Leave Requests", "/leave-requests", "Leave workflow and proxy relevance"),
        ("Exam Schedule", "/exam-schedule", "Exam schedule builder and generated grid"),
        ("Settings", "/settings", "Institution and user preference management"),
        ("Search and Notifications", "/search/features, /notifications", "JSON endpoints for global navigation and notification center"),
    ], (1.45, 2.8, 2.25))

    h(doc, "Appendix C: Test Evidence and Submission Checklist", 1)
    h(doc, "C.1 Command Evidence", 2)
    table(doc, [
        ("Command", "Result", "Notes"),
        ("cmd /c npm run build", "Passed", "Vite production build completed successfully."),
        ("php artisan route:list", "Passed", "79 routes available in the local application."),
        ("php artisan test", "43 passing, 3 failing", "The failing tests are stale default auth/root tests that need updating for the Firebase-aware flow and redirecting root route."),
        ("DOCX package scan", "Passed", "No old brand text and no em dash or en dash characters were found."),
        ("Header/footer scan", "Passed", "Word header and footer XML parts contain no visible text."),
    ], (2.1, 1.55, 2.85))
    h(doc, "C.2 Final Artifact Checklist", 2)
    table(doc, [
        ("Required Artifact", "Included in Document"),
        ("Cover page", "Yes"),
        ("Acknowledgement", "Yes"),
        ("Table of contents", "Yes"),
        ("Introduction and objectives", "Yes"),
        ("Functional and non-functional requirements", "Yes"),
        ("Use case diagram", "Yes"),
        ("Data flow diagram", "Yes"),
        ("Architecture diagram", "Yes"),
        ("Entity relationship diagram", "Yes"),
        ("Database schema and data dictionary", "Yes"),
        ("Route and API design", "Yes"),
        ("Frontend component structure", "Yes"),
        ("Authentication sequence diagram", "Yes"),
        ("Proxy workflow diagram", "Yes"),
        ("UI mockups and screen result table", "Yes"),
        ("Testing approach and test cases", "Yes"),
        ("Limitations, future work, conclusion, references", "Yes"),
    ], (3.2, 2.4))


def build():
    logo = ASSETS / "scholarly-logo.png"
    draw_logo(logo)
    draw_architecture(ASSETS / "architecture.png")
    draw_usecase(ASSETS / "usecase.png")
    draw_erd(ASSETS / "erd.png")
    draw_results(ASSETS / "results.png")
    draw_data_flow(ASSETS / "data-flow.png")
    draw_auth_sequence(ASSETS / "auth-sequence.png")
    draw_proxy_workflow(ASSETS / "proxy-workflow.png")
    draw_component_map(ASSETS / "component-map.png")
    draw_class_model(ASSETS / "model-relationships.png")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    cover_logo = doc.add_paragraph()
    cover_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_logo.add_run().add_picture(str(logo), width=Inches(1.25))
    p(doc, "Scholarly", bold=True, color=DARK, size=34, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    p(doc, "Institution Management and Classroom Operations System", color=GREEN_DARK, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    p(doc, "Project Documentation", bold=True, color=INK, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    p(doc, "Full-stack web application built with Laravel, Vue.js, Inertia, MySQL, Firebase, and WhatsApp Cloud API integration.", color=MUTED, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    kv(doc, [
        ("Project Name", "Scholarly"),
        ("Project Type", "Academic full-stack web application"),
        ("Prepared By", "Shakif Niaz"),
        ("Student ID", "To be added"),
        ("Supervisor", "To be added"),
        ("Department", "Software Engineering Department"),
        ("Submission Date", "August 2026"),
        ("Development Status", "Working local prototype with seeded institute data"),
    ], widths=(2.0, 4.3))
    p(doc, "This document follows Scholarly's green and deep-ink interface palette. No running headers or footers are used.", color=MUTED, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, before=18)
    doc.add_page_break()

    add_long_content(doc)

    for section in doc.sections:
        for part in [section.header, section.footer, section.first_page_header, section.first_page_footer, section.even_page_header, section.even_page_footer]:
            for paragraph in part.paragraphs:
                paragraph.text = ""

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
                continue
            run.text = clean(run.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
                            continue
                        run.text = clean(run.text)

    props = doc.core_properties
    props.title = "Scholarly Project Documentation"
    props.subject = "Laravel and Vue.js school operations system documentation"
    props.author = "Shakif Niaz"
    props.comments = "Generated project documentation for Scholarly."
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
